import os
from datetime import datetime
import json
import logging
import sqlite3
from pathlib import Path
from uuid import uuid4
import tempfile
import threading
import queue

from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles

from db import get_db, init_db, upsert_councillor, list_councillors
from xml_utils import parse_agenda_xml
from ai_utils import align_questions_with_vtt

from docx import Document
from typing import Optional, List, Dict

from dotenv import load_dotenv
load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("quest")

app = FastAPI()

static_dir = Path(__file__).parent / "static"
app.mount("/static", StaticFiles(directory=static_dir, html=True), name="static")


def _default_storage_dir() -> Path:
    db_path = os.environ.get("QUEST_DB_PATH")
    if db_path:
        return Path(db_path).resolve().parent / "uploads"
    return Path(__file__).parent / "data" / "uploads"


storage_dir = Path(os.environ.get("QUEST_STORAGE_DIR", _default_storage_dir()))
storage_dir.mkdir(parents=True, exist_ok=True)


def _now_iso() -> str:
    return datetime.utcnow().isoformat()


def _resolve_source_question(meeting_data: Dict, question_data: Dict) -> Dict:
    """Zoek de oorspronkelijke vraag uit de opgeslagen XML."""
    raw_source = meeting_data.get("source_questions_json") or ""
    parsed_source = []
    if raw_source:
        try:
            parsed_source = json.loads(raw_source)
        except json.JSONDecodeError:
            parsed_source = []
    candidate = None
    idx = question_data.get("source_question_idx")
    if isinstance(idx, int) and 0 <= idx < len(parsed_source):
        candidate = parsed_source[idx]
    if not candidate:
        dossier = (question_data.get("dossier_id") or "").strip()
        seq = (question_data.get("sequence_nr") or "").strip()
        for item in parsed_source:
            if dossier and (item.get("dossier_id") or "").strip() == dossier:
                candidate = item
                break
            if not dossier and seq and (item.get("sequence_nr") or "").strip() == seq:
                candidate = item
                break
    if not candidate:
        candidate = {
            "meeting_date": meeting_data.get("meeting_date"),
            "commission_name": meeting_data.get("commission_name"),
            "dossier_id": question_data.get("dossier_id"),
            "dossier_year_nr": question_data.get("dossier_year_nr"),
            "sequence_nr": question_data.get("sequence_nr"),
            "title": question_data.get("title"),
            "subject": question_data.get("subject"),
            "roi_type": question_data.get("roi_type"),
            "submitter_given_name": question_data.get("submitter_given_name"),
            "submitter_family_name": question_data.get("submitter_family_name"),
            "submitter_faction": question_data.get("submitter_faction"),
            "assignee_label": question_data.get("assignee_label"),
            "assignee_given_name": question_data.get("assignee_given_name"),
            "assignee_family_name": question_data.get("assignee_family_name"),
            "question_text_from_xml": "",
        }
    candidate = dict(candidate)
    candidate.setdefault("meeting_date", meeting_data.get("meeting_date"))
    candidate.setdefault("commission_name", meeting_data.get("commission_name"))
    question_text = (
        candidate.get("question_text_from_xml")
        or question_data.get("question_text_xml")
        or question_data.get("question_text_raw")
        or ""
    )
    candidate["question_text_from_xml"] = question_text
    return candidate


def _update_meeting_processing_summary(meeting_id: int, manual_error: str = "") -> Dict:
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT processing_state, COUNT(*) AS amount
        FROM questions
        WHERE meeting_id = ?
        GROUP BY processing_state
        """,
        (meeting_id,),
    )
    counts = {row["processing_state"] or "": row["amount"] for row in cur.fetchall()}
    total = sum(counts.values())
    completed = counts.get("completed", 0)
    queued = counts.get("queued", 0) + counts.get("pending", 0)
    running = counts.get("in_progress", 0)
    errors = counts.get("error", 0)
    if total == 0:
        state = "completed"
    elif queued or running:
        state = "processing"
    elif errors:
        state = "error"
    else:
        state = "completed"
    message = manual_error or (
        f"{errors} vraag{'en' if errors != 1 else ''} met fout."
        if errors
        else ""
    )
    complete_ts = _now_iso()
    cur.execute(
        """
        UPDATE meetings
        SET processed_questions = ?,
            total_questions = ?,
            processing_state = ?,
            processing_completed_at = CASE WHEN ? = 'completed' THEN ? ELSE processing_completed_at END,
            processing_error = ?
        WHERE id = ?
        """,
        (completed, total, state, state, complete_ts, message, meeting_id),
    )
    conn.commit()
    conn.close()
    return {
        "state": state,
        "total": total,
        "completed": completed,
        "queued": queued,
        "running": running,
        "errors": errors,
        "message": message,
    }


class QuestionProcessingQueue:
    def __init__(self):
        self._queue: "queue.Queue[Optional[int]]" = queue.Queue()
        self._stop_event = threading.Event()
        self._thread = threading.Thread(target=self._worker, name="question-processor", daemon=True)
        self._started = False

    def start(self):
        if self._started:
            return
        self._started = True
        self._restore_pending_jobs()
        self._thread.start()
        logger.info("QuestionProcessingQueue started (pending items restored).")

    def stop(self):
        self._stop_event.set()
        self._queue.put(None)
        if self._thread.is_alive():
            self._thread.join(timeout=5)
        logger.info("QuestionProcessingQueue stopped.")

    def _restore_pending_jobs(self):
        conn = get_db()
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id
            FROM questions
            WHERE processing_state IN ('pending', 'queued', 'in_progress')
            ORDER BY id
            """
        )
        ids = [row["id"] for row in cur.fetchall()]
        for question_id in ids:
            self._queue.put(question_id)
        conn.close()
        if ids:
            logger.info("Restored %d vragen naar de queue na herstart.", len(ids))

    def enqueue_meeting(self, meeting_id: int):
        conn = get_db()
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id
            FROM questions
            WHERE meeting_id = ?
              AND processing_state IN ('pending', 'error')
            ORDER BY id
            """,
            (meeting_id,),
        )
        ids = [row["id"] for row in cur.fetchall()]
        if ids:
            cur.executemany(
                "UPDATE questions SET processing_state = 'queued', processing_error = '' WHERE id = ?",
                [(qid,) for qid in ids],
            )
        cur.execute(
            """
            UPDATE meetings
            SET processing_state = CASE
                    WHEN total_questions > 0 THEN 'queued'
                    ELSE processing_state
                END,
                processing_started_at = CASE
                    WHEN processing_started_at IS NULL AND total_questions > 0 THEN ?
                    ELSE processing_started_at
                END
            WHERE id = ?
            """,
            (_now_iso(), meeting_id),
        )
        conn.commit()
        conn.close()
        for question_id in ids:
            self._queue.put(question_id)

    def enqueue_question(self, question_id: int):
        conn = get_db()
        cur = conn.cursor()
        cur.execute(
            "UPDATE questions SET processing_state = 'queued', processing_error = '' WHERE id = ?",
            (question_id,),
        )
        conn.commit()
        conn.close()
        self._queue.put(question_id)

    def stats(self) -> Dict:
        return {"queued_in_memory": self._queue.qsize()}

    def _worker(self):
        while not self._stop_event.is_set():
            try:
                question_id = self._queue.get(timeout=1.0)
            except queue.Empty:
                continue
            if question_id is None:
                break
            try:
                self._process_question(question_id)
            except Exception:
                logger.exception("Onverwachte fout tijdens verwerken van vraag %s", question_id)
            finally:
                self._queue.task_done()

    def _process_question(self, question_id: int):
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT * FROM questions WHERE id = ?", (question_id,))
        question = cur.fetchone()
        if not question:
            conn.close()
            return
        if (question["processing_state"] or "").lower() == "completed":
            conn.close()
            return
        cur.execute("SELECT * FROM meetings WHERE id = ?", (question["meeting_id"],))
        meeting = cur.fetchone()
        if not meeting:
            conn.close()
            return
        start_ts = _now_iso()
        cur.execute(
            """
            UPDATE questions
            SET processing_state = 'in_progress',
                processing_started_at = COALESCE(processing_started_at, ?),
                processing_error = ''
            WHERE id = ?
            """,
            (start_ts, question_id),
        )
        cur.execute(
            """
            UPDATE meetings
            SET processing_state = 'processing',
                processing_started_at = COALESCE(processing_started_at, ?)
            WHERE id = ?
            """,
            (start_ts, meeting["id"]),
        )
        conn.commit()
        councillors = list_councillors(conn)
        meeting_data = dict(meeting)
        question_data = dict(question)
        transcript_text = meeting_data.get("transcript_text") or ""
        conn.close()

        if not transcript_text.strip():
            self._mark_question_error(question_id, meeting_data["id"], "Geen transcript beschikbaar voor deze vergadering.")
            return

        source_question = _resolve_source_question(meeting_data, question_data)
        try:
            ai_items = align_questions_with_vtt([source_question], transcript_text, councillors)
        except Exception as exc:
            self._mark_question_error(question_id, meeting_data["id"], str(exc))
            return

        if not ai_items:
            self._mark_question_error(question_id, meeting_data["id"], "Geen AI-resultaat ontvangen.")
            return

        item = ai_items[0]
        item.setdefault("answer_text_verbatim", item.get("answer_text_raw", ""))
        item.setdefault("answer_text_raw", "")
        item.setdefault("answer_status", "draft")
        finish_ts = _now_iso()

        conn = get_db()
        cur = conn.cursor()
        cur.execute(
            """
            UPDATE questions
            SET
                question_start_time = ?,
                question_end_time = ?,
                answer_start_time = ?,
                answer_end_time = ?,
                question_text_raw = ?,
                answer_text_verbatim = ?,
                answer_text_raw = ?,
                summary = ?,
                actions_json = ?,
                topics_json = ?,
                note = ?,
                answer_status = ?,
                processing_state = 'completed',
                processing_completed_at = ?,
                processing_error = '',
                processing_attempts = processing_attempts + 1
            WHERE id = ?
            """,
            (
                item.get("question_start_time") or "",
                item.get("question_end_time") or "",
                item.get("answer_start_time") or "",
                item.get("answer_end_time") or "",
                item.get("question_text_raw") or question_data.get("question_text_raw") or "",
                item.get("answer_text_verbatim") or "",
                item.get("answer_text_raw") or "",
                item.get("summary") or "",
                json.dumps(item.get("actions") or [], ensure_ascii=False),
                json.dumps(item.get("topics") or [], ensure_ascii=False),
                item.get("note") or "",
                item.get("answer_status") or "draft",
                finish_ts,
                question_id,
            ),
        )
        conn.commit()
        conn.close()
        _update_meeting_processing_summary(meeting_data["id"])

    def _mark_question_error(self, question_id: int, meeting_id: int, message: str):
        truncated = (message or "")[:500]
        conn = get_db()
        cur = conn.cursor()
        cur.execute(
            """
            UPDATE questions
            SET processing_state = 'error',
                processing_error = ?,
                processing_completed_at = NULL,
                processing_attempts = processing_attempts + 1
            WHERE id = ?
            """,
            (truncated, question_id),
        )
        conn.commit()
        conn.close()
        _update_meeting_processing_summary(meeting_id, manual_error=truncated)


processing_queue: Optional[QuestionProcessingQueue] = None


def _enqueue_meeting_processing(meeting_id: int):
    if processing_queue:
        processing_queue.enqueue_meeting(meeting_id)


def _enqueue_question_processing(question_id: int):
    if processing_queue:
        processing_queue.enqueue_question(question_id)


def _coerce_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
        except json.JSONDecodeError:
            parsed = None
        if isinstance(parsed, list):
            return [str(v).strip() for v in parsed if str(v).strip()]
        return [part.strip() for part in value.split(",") if part.strip()]
    return []


def _deserialize_question_row(row):
    data = dict(row)
    data["actions"] = _coerce_list(data.get("actions_json"))
    data["topics"] = _coerce_list(data.get("topics_json"))
    return data


QUESTION_INSERT_COLUMNS = [
    "meeting_id",
    "dossier_id",
    "dossier_year_nr",
    "sequence_nr",
    "title",
    "subject",
    "roi_type",
    "submitter_given_name",
    "submitter_family_name",
    "submitter_faction",
    "assignee_label",
    "assignee_given_name",
    "assignee_family_name",
    "question_start_time",
    "question_end_time",
    "answer_start_time",
    "answer_end_time",
    "reply_start_time",
    "reply_end_time",
    "question_text_raw",
    "answer_text_verbatim",
    "answer_text_raw",
    "question_text_xml",
    "summary",
    "actions_json",
    "topics_json",
    "note",
    "answer_status",
    "processing_state",
    "processing_error",
    "processing_started_at",
    "processing_completed_at",
    "processing_attempts",
    "source_question_idx",
]

QUESTION_INSERT_SQL = (
    "INSERT INTO questions ("
    + ", ".join(QUESTION_INSERT_COLUMNS)
    + ") VALUES ("
    + ",".join(["?"] * len(QUESTION_INSERT_COLUMNS))
    + ")"
)


@app.on_event("startup")
def startup_event():
    global processing_queue
    init_db()
    if processing_queue is None:
        processing_queue = QuestionProcessingQueue()
    processing_queue.start()


@app.on_event("shutdown")
def shutdown_event():
    if processing_queue:
        processing_queue.stop()


@app.get("/", response_class=HTMLResponse)
def index():
    index_html = (static_dir / "index.html").read_text(encoding="utf-8")
    return index_html


@app.get("/meeting/{meeting_id}", response_class=HTMLResponse)
def meeting_page(meeting_id: int):
    html = (static_dir / "meeting.html").read_text(encoding="utf-8")
    return html


@app.get("/questions", response_class=HTMLResponse)
def questions_page():
    html = (static_dir / "questions.html").read_text(encoding="utf-8")
    return html


@app.get("/councillors", response_class=HTMLResponse)
def councillors_page():
    html = (static_dir / "councillors.html").read_text(encoding="utf-8")
    return html


def _sanitize_filename(name: str, fallback: str) -> str:
    candidate = (name or fallback or "uploaded").strip().lower()
    allowed = [c if c.isalnum() else "-" for c in candidate]
    sanitized = "".join(allowed).strip("-")
    return sanitized or fallback or "uploaded"


@app.post("/api/upload")
async def upload(
    agenda: UploadFile = File(...),
    transcript: UploadFile = File(...),
    webcast_id: str = Form(""),
):
    logger.info(
        "Upload started webcast_id=%s agenda=%s transcript=%s",
        webcast_id,
        getattr(agenda, "filename", "unknown"),
        getattr(transcript, "filename", "unknown"),
    )
    agenda_bytes = await agenda.read()
    transcript_bytes = await transcript.read()
    xml_str = agenda_bytes.decode("utf-8", errors="ignore")
    vtt_str = transcript_bytes.decode("utf-8", errors="ignore")

    upload_dir = storage_dir / (
        datetime.utcnow().strftime("%Y%m%d-%H%M%S") + f"-{uuid4().hex[:8]}"
    )
    upload_dir.mkdir(parents=True, exist_ok=True)
    agenda_path = upload_dir / f"agenda-{_sanitize_filename(agenda.filename, 'xml')}.xml"
    transcript_path = upload_dir / (
        f"transcript-{_sanitize_filename(transcript.filename, 'vtt')}.vtt"
    )
    agenda_path.write_bytes(agenda_bytes)
    transcript_path.write_bytes(transcript_bytes)

    oral_questions, meeting_date, commission_name = parse_agenda_xml(xml_str)
    question_lookup = {
        q.get("dossier_id") or f"idx-{idx}": q
        for idx, q in enumerate(oral_questions)
    }
    conn = get_db()

    def register_person(given: str, family: str, titled: str):
        upsert_councillor(conn, given, family, titled or "")

    for q in oral_questions:
        submitter_full_title = ""
        if q.get("submitter_given_name") or q.get("submitter_family_name"):
            submitter_full_title = "raadslid " + " ".join(
                part for part in (q.get("submitter_given_name"), q.get("submitter_family_name")) if part
            )
        register_person(
            q.get("submitter_given_name", ""),
            q.get("submitter_family_name", ""),
            submitter_full_title.strip(),
        )
        assignee_title = q.get("assignee_label") or ""
        register_person(
            q.get("assignee_given_name", ""),
            q.get("assignee_family_name", ""),
            assignee_title,
        )

    conn.commit()
    logger.info(
        "Parsed XML -> %d questions (meeting_date=%s, commission=%s)",
        len(oral_questions),
        meeting_date,
        commission_name,
    )

    cur = conn.cursor()
    processing_started_at = _now_iso()
    initial_state = "queued" if oral_questions else "completed"
    cur.execute(
        """INSERT INTO meetings (
            meeting_date, commission_name, webcast_id,
            source_questions_json, transcript_text,
            agenda_file_path, transcript_file_path,
            processing_state, processing_started_at,
            total_questions, processed_questions, processing_error
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (
            meeting_date,
            commission_name,
            webcast_id,
            json.dumps(oral_questions, ensure_ascii=False),
            vtt_str,
            str(agenda_path),
            str(transcript_path),
            initial_state,
            processing_started_at if oral_questions else None,
            len(oral_questions),
            0,
            "",
        ),
    )
    meeting_id = cur.lastrowid
    logger.info("Stored meeting id=%s", meeting_id)

    for idx, q in enumerate(oral_questions):
        key = q.get("dossier_id") or f"idx-{idx}"
        cur.execute(
            QUESTION_INSERT_SQL,
            (
                meeting_id,
                q.get("dossier_id"),
                q.get("dossier_year_nr"),
                q.get("sequence_nr"),
                q.get("title"),
                q.get("subject"),
                q.get("roi_type"),
                q.get("submitter_given_name"),
                q.get("submitter_family_name"),
                q.get("submitter_faction"),
                q.get("assignee_label"),
                q.get("assignee_given_name"),
                q.get("assignee_family_name"),
                "",
                "",
                "",
                "",
                "",
                "",
                q.get("question_text_from_xml", ""),
                "",
                "",
                q.get("question_text_from_xml", ""),
                "",
                json.dumps([], ensure_ascii=False),
                json.dumps([], ensure_ascii=False),
                "Ingeladen vanuit XML, wacht op verwerking.",
                "draft",
                "pending",
                "",
                None,
                None,
                0,
                idx,
            ),
        )
        logger.info(
            "Stored placeholder for question sequence_nr=%s (db id=%s)",
            q.get("sequence_nr"),
            cur.lastrowid,
        )

    conn.commit()
    conn.close()

    if oral_questions:
        _enqueue_meeting_processing(meeting_id)

    final_status = "completed" if not oral_questions else "queued"
    return {
        "status": final_status,
        "meeting_id": meeting_id,
        "questions": len(oral_questions),
    }


@app.get("/api/meetings/{meeting_id}")
def get_meeting(meeting_id: int):
    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT * FROM meetings WHERE id = ?", (meeting_id,))
    meeting = cur.fetchone()
    if not meeting:
        conn.close()
        return JSONResponse({"error": "meeting not found"}, status_code=404)

    cur.execute(
        """
        SELECT *
        FROM questions
        WHERE meeting_id = ?
        ORDER BY
            CASE WHEN question_start_time IS NULL OR question_start_time = '' THEN 1 ELSE 0 END,
            question_start_time,
            sequence_nr
        """,
        (meeting_id,),
    )
    questions = [_deserialize_question_row(row) for row in cur.fetchall()]
    conn.close()

    return {"meeting": dict(meeting), "questions": questions}


from pydantic import BaseModel


class QuestionCreate(BaseModel):
    meeting_id: int
    dossier_id: Optional[str] = None
    dossier_year_nr: Optional[str] = None
    sequence_nr: Optional[str] = None
    title: Optional[str] = ""
    subject: Optional[str] = ""
    roi_type: Optional[str] = ""
    submitter_given_name: Optional[str] = ""
    submitter_family_name: Optional[str] = ""
    submitter_faction: Optional[str] = ""
    assignee_label: Optional[str] = ""
    assignee_given_name: Optional[str] = ""
    assignee_family_name: Optional[str] = ""
    question_start_time: Optional[str] = ""
    question_end_time: Optional[str] = ""
    answer_start_time: Optional[str] = ""
    answer_end_time: Optional[str] = ""
    question_text_raw: Optional[str] = ""
    answer_text_verbatim: Optional[str] = ""
    answer_text_raw: Optional[str] = ""


class QuestionUpdate(BaseModel):
    question_start_time: Optional[str] = None
    question_end_time: Optional[str] = None
    answer_start_time: Optional[str] = None
    answer_end_time: Optional[str] = None
    reply_start_time: Optional[str] = None
    reply_end_time: Optional[str] = None
    answer_text_verbatim: Optional[str] = None
    answer_text_raw: Optional[str] = None
    answer_status: Optional[str] = None
    summary: Optional[str] = None
    actions: Optional[List[str]] = None
    topics: Optional[List[str]] = None


@app.post("/api/questions", status_code=201)
async def create_question(payload: QuestionCreate):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT id FROM meetings WHERE id = ?", (payload.meeting_id,))
    meeting = cur.fetchone()
    if not meeting:
        conn.close()
        return JSONResponse({"error": "meeting not found"}, status_code=404)

    sequence_nr = payload.sequence_nr
    if not sequence_nr:
        cur.execute("SELECT sequence_nr FROM questions WHERE meeting_id = ?", (payload.meeting_id,))
        rows = [row[0] for row in cur.fetchall()]
        max_seq = 0
        for value in rows:
            try:
                num = int(str(value))
            except (ValueError, TypeError):
                continue
            if num > max_seq:
                max_seq = num
        sequence_nr = str(max_seq + 1)

    now_ts = datetime.utcnow().isoformat()
    cur.execute(
        QUESTION_INSERT_SQL,
        (
            payload.meeting_id,
            payload.dossier_id,
            payload.dossier_year_nr,
            sequence_nr,
            (payload.title or payload.subject or "").strip(),
            (payload.subject or payload.title or "").strip(),
            (payload.roi_type or "").strip(),
            (payload.submitter_given_name or "").strip(),
            (payload.submitter_family_name or "").strip(),
            (payload.submitter_faction or "").strip(),
            (payload.assignee_label or "").strip(),
            (payload.assignee_given_name or "").strip(),
            (payload.assignee_family_name or "").strip(),
            (payload.question_start_time or "").strip(),
            (payload.question_end_time or "").strip(),
            (payload.answer_start_time or "").strip(),
            (payload.answer_end_time or "").strip(),
            "",
            "",
            (payload.question_text_raw or payload.subject or "").strip(),
            (payload.answer_text_verbatim or payload.answer_text_raw or "").strip(),
            (payload.answer_text_raw or "").strip(),
            "",
            "",
            json.dumps([], ensure_ascii=False),
            json.dumps([], ensure_ascii=False),
            "Handmatig toegevoegd via interface.",
            "draft",
            "completed",
            "",
            now_ts,
            now_ts,
            0,
            None,
        ),
    )
    question_id = cur.lastrowid
    conn.commit()
    cur.execute("SELECT * FROM questions WHERE id = ?", (question_id,))
    question = cur.fetchone()
    conn.close()
    _update_meeting_processing_summary(payload.meeting_id)
    return {"status": "created", "question": _deserialize_question_row(question)}


class CouncillorCreate(BaseModel):
    given_name: str = ""
    family_name: str = ""
    name_with_title: Optional[str] = ""
    wrong_spellings: Optional[str] = ""


class CouncillorUpdate(BaseModel):
    given_name: Optional[str] = None
    family_name: Optional[str] = None
    name_with_title: Optional[str] = None
    wrong_spellings: Optional[str] = None


@app.patch("/api/questions/{question_id}")
async def update_question(question_id: int, payload: QuestionUpdate):
    data = payload.model_dump(exclude_unset=True)
    if not data:
        return {"status": "no changes"}

    allowed_status = {"draft", "approved"}
    normalized = {}
    for key, value in data.items():
        if key in {"answer_text_raw", "answer_text_verbatim"}:
            normalized[key] = value or ""
        elif key == "answer_status":
            status = (value or "").strip().lower() or "draft"
            if status not in allowed_status:
                return JSONResponse(
                    {"error": f"answer_status moet 'draft' of 'approved' zijn (kreeg: {value})"},
                    status_code=400,
                )
            normalized[key] = status
        elif key in {"actions", "topics"}:
            normalized[f"{key}_json"] = json.dumps(_coerce_list(value), ensure_ascii=False)
        elif key == "summary":
            normalized[key] = value or ""
        else:
            normalized[key] = value

    conn = get_db()
    cur = conn.cursor()

    fields = []
    values = []
    for k, v in normalized.items():
        fields.append(f"{k} = ?")
        values.append(v)

    values.append(question_id)
    sql = f"UPDATE questions SET {', '.join(fields)} WHERE id = ?"
    cur.execute(sql, values)
    conn.commit()
    conn.close()

    return {"status": "ok"}


@app.delete("/api/questions/{question_id}")
def delete_question(question_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT meeting_id FROM questions WHERE id = ?", (question_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        return JSONResponse({"error": "question not found"}, status_code=404)

    cur.execute("DELETE FROM questions WHERE id = ?", (question_id,))
    conn.commit()
    conn.close()
    _update_meeting_processing_summary(row["meeting_id"])
    return {"status": "deleted", "question_id": question_id}


@app.post("/api/questions/{question_id}/regenerate")
async def regenerate_question(question_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM questions WHERE id = ?", (question_id,))
    question = cur.fetchone()
    if not question:
        conn.close()
        return JSONResponse({"error": "question not found"}, status_code=404)

    cur.execute("SELECT * FROM meetings WHERE id = ?", (question["meeting_id"],))
    meeting = cur.fetchone()
    if not meeting:
        conn.close()
        return JSONResponse({"error": "meeting not found"}, status_code=404)

    transcript_text = meeting["transcript_text"]
    if not transcript_text:
        conn.close()
        return JSONResponse(
            {
                "error": "Voor deze vergadering is geen transcript opgeslagen. Upload de vergadering opnieuw om dit mogelijk te maken."
            },
            status_code=400,
        )

    cur.execute(
        """
        UPDATE questions
        SET processing_state = 'pending',
            processing_error = '',
            processing_started_at = NULL,
            processing_completed_at = NULL
        WHERE id = ?
        """,
        (question_id,),
    )
    conn.commit()
    cur.execute("SELECT * FROM questions WHERE id = ?", (question_id,))
    updated = cur.fetchone()
    conn.close()
    _update_meeting_processing_summary(question["meeting_id"])
    _enqueue_question_processing(question_id)
    return {"status": "queued", "question": _deserialize_question_row(updated)}


@app.get("/export/{meeting_id}")
def export_docx(meeting_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM meetings WHERE id = ?", (meeting_id,))
    meeting = cur.fetchone()
    if not meeting:
        conn.close()
        return JSONResponse({"error": "meeting not found"}, status_code=404)

    cur.execute(
        """
        SELECT *
        FROM questions
        WHERE meeting_id = ?
        ORDER BY
            CASE WHEN question_start_time IS NULL OR question_start_time = '' THEN 1 ELSE 0 END,
            question_start_time,
            sequence_nr
        """,
        (meeting_id,),
    )
    questions = cur.fetchall()
    conn.close()

    doc = Document()
    doc.add_heading(f"{meeting['commission_name']} - {meeting['meeting_date']}", level=1)

    for q in questions:
        doc.add_heading(f"{q['sequence_nr']} - {q['title']}", level=2)
        doc.add_paragraph(f"Dossier: {q['dossier_year_nr']}")
        doc.add_paragraph(
            f"Vraagsteller: {q['submitter_given_name']} {q['submitter_family_name']} ({q['submitter_faction']})"
        )
        doc.add_paragraph(f"Bevoegde schepen: {q['assignee_label']}")
        doc.add_paragraph("\nVraag:")
        doc.add_paragraph(q['question_text_raw'] or "")
        doc.add_paragraph("\nAntwoord (quasi letterlijke versie):")
        doc.add_paragraph(q['answer_text_verbatim'] or "")
        doc.add_paragraph("\nAntwoord (gebalde versie):")
        doc.add_paragraph(q['answer_text_raw'] or "")
        status = (q['answer_status'] or "draft").capitalize()
        doc.add_paragraph(f"Status: {status}")
        doc.add_page_break()

    tmp = Path(tempfile.gettempdir()) / f"vragen_{meeting_id}.docx"
    doc.save(tmp)
    return FileResponse(tmp, filename=tmp.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.delete("/api/meetings/{meeting_id}")
def delete_meeting(meeting_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT id FROM meetings WHERE id = ?", (meeting_id,))
    if not cur.fetchone():
        conn.close()
        return JSONResponse({"error": "meeting not found"}, status_code=404)

    cur.execute("DELETE FROM questions WHERE meeting_id = ?", (meeting_id,))
    cur.execute("DELETE FROM meetings WHERE id = ?", (meeting_id,))
    conn.commit()
    conn.close()
    return {"status": "deleted", "meeting_id": meeting_id}


@app.post("/api/meetings/{meeting_id}/restore-missing")
def restore_missing_questions(meeting_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM meetings WHERE id = ?", (meeting_id,))
    meeting = cur.fetchone()
    if not meeting:
        conn.close()
        return JSONResponse({"error": "meeting not found"}, status_code=404)

    source_json = meeting["source_questions_json"]
    if not source_json:
        conn.close()
        return JSONResponse({"error": "no source questions stored for this meeting"}, status_code=400)

    oral_questions = json.loads(source_json)
    def question_key(data: dict):
        return data.get("dossier_id") or f"seq-{data.get('sequence_nr')}"

    cur.execute("SELECT dossier_id, sequence_nr FROM questions WHERE meeting_id = ?", (meeting_id,))
    existing_ids = {
        row["dossier_id"] or f"seq-{row['sequence_nr']}"
        for row in cur.fetchall()
        if row["dossier_id"] or row["sequence_nr"]
    }

    added = 0
    for idx, q in enumerate(oral_questions):
        key = question_key(q)
        if key in existing_ids:
            continue

        cur.execute(
            QUESTION_INSERT_SQL,
            (
                meeting_id,
                q.get("dossier_id"),
                q.get("dossier_year_nr"),
                q.get("sequence_nr"),
                q.get("title"),
                q.get("subject"),
                q.get("roi_type"),
                q.get("submitter_given_name"),
                q.get("submitter_family_name"),
                q.get("submitter_faction"),
                q.get("assignee_label"),
                q.get("assignee_given_name"),
                q.get("assignee_family_name"),
                "",
                "",
                "",
                "",
                "",
                "",
                q.get("question_text_from_xml", ""),
                "",
                "",
                q.get("question_text_from_xml", ""),
                "",
                json.dumps([], ensure_ascii=False),
                json.dumps([], ensure_ascii=False),
                "Automatisch toegevoegd vanuit bron-XML (geen AI-resultaat).",
                "draft",
                "pending",
                "",
                None,
                None,
                0,
                idx,
            ),
        )
        added += 1
        existing_ids.add(key)

    conn.commit()
    conn.close()
    if added:
        _enqueue_meeting_processing(meeting_id)
    _update_meeting_processing_summary(meeting_id)
    return {"status": "ok", "added": added}


@app.get("/api/meetings")
def list_meetings():
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT m.*, COUNT(q.id) AS question_count
        FROM meetings m
        LEFT JOIN questions q ON q.meeting_id = m.id
        GROUP BY m.id
        ORDER BY m.meeting_date DESC, m.id DESC
        """
    )
    meetings = [dict(row) for row in cur.fetchall()]
    conn.close()
    return {"meetings": meetings}


@app.get("/api/processing/meetings/{meeting_id}")
def get_meeting_processing_status(meeting_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT id FROM meetings WHERE id = ?", (meeting_id,))
    meeting_exists = cur.fetchone()
    conn.close()
    if not meeting_exists:
        return JSONResponse({"error": "meeting not found"}, status_code=404)

    summary = _update_meeting_processing_summary(meeting_id)
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        "SELECT processing_state, total_questions, processed_questions, processing_error FROM meetings WHERE id = ?",
        (meeting_id,),
    )
    meeting_row = cur.fetchone()
    cur.execute(
        """
        SELECT processing_state, COUNT(*) AS amount
        FROM questions
        WHERE meeting_id = ?
        GROUP BY processing_state
        """,
        (meeting_id,),
    )
    states = {row["processing_state"] or "": row["amount"] for row in cur.fetchall()}
    conn.close()
    summary.update(
        {
            "meeting_id": meeting_id,
            "processing_state": meeting_row["processing_state"],
            "processed_questions": meeting_row["processed_questions"],
            "total_questions": meeting_row["total_questions"],
            "processing_error": meeting_row["processing_error"],
            "states": states,
        }
    )
    return summary


@app.get("/api/processing/queue")
def get_queue_status():
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT
            m.id AS meeting_id,
            m.meeting_date,
            m.commission_name,
            SUM(CASE WHEN q.processing_state IN ('pending', 'queued') THEN 1 ELSE 0 END) AS queued_questions,
            SUM(CASE WHEN q.processing_state = 'in_progress' THEN 1 ELSE 0 END) AS in_progress_questions,
            SUM(CASE WHEN q.processing_state = 'error' THEN 1 ELSE 0 END) AS error_questions,
            COUNT(q.id) AS total_questions
        FROM meetings m
        JOIN questions q ON q.meeting_id = m.id
        GROUP BY m.id
        HAVING
            queued_questions > 0
            OR in_progress_questions > 0
            OR error_questions > 0
        ORDER BY m.meeting_date DESC, m.id DESC
        """
    )
    meetings = []
    for row in cur.fetchall():
        meetings.append(
            {
                "meeting_id": row["meeting_id"],
                "meeting_date": row["meeting_date"],
                "commission_name": row["commission_name"],
                "queued": row["queued_questions"],
                "in_progress": row["in_progress_questions"],
                "errors": row["error_questions"],
                "total": row["total_questions"],
            }
        )
    conn.close()
    queue_snapshot = processing_queue.stats() if processing_queue else {"queued_in_memory": 0}
    return {"meetings": meetings, "queue": queue_snapshot}


@app.get("/api/councillors")
def get_councillors():
    conn = get_db()
    data = list_councillors(conn)
    conn.close()
    return {"councillors": data}


@app.get("/api/question-people")
def get_question_people():
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT DISTINCT
          TRIM(
            TRIM(COALESCE(submitter_given_name, '')) || ' ' || TRIM(COALESCE(submitter_family_name, ''))
          ) AS full_name
        FROM questions
        WHERE TRIM(
          COALESCE(submitter_given_name, '') || COALESCE(submitter_family_name, '')
        ) != ''
        ORDER BY LOWER(full_name)
        """
    )
    submitters = [row["full_name"] for row in cur.fetchall() if row["full_name"]]

    cur.execute(
        """
        SELECT DISTINCT
          TRIM(
            CASE
              WHEN LENGTH(TRIM(COALESCE(assignee_label, ''))) > 0 THEN assignee_label
              ELSE TRIM(COALESCE(assignee_given_name, '') || ' ' || COALESCE(assignee_family_name, ''))
            END
          ) AS full_name
        FROM questions
        WHERE TRIM(
          COALESCE(assignee_label, '') ||
          COALESCE(assignee_given_name, '') ||
          COALESCE(assignee_family_name, '')
        ) != ''
        ORDER BY LOWER(full_name)
        """
    )
    assignees = [row["full_name"] for row in cur.fetchall() if row["full_name"]]

    cur.execute("SELECT topics_json FROM questions WHERE topics_json IS NOT NULL AND topics_json != ''")
    topic_set = set()
    for row in cur.fetchall():
        try:
            items = json.loads(row["topics_json"])
        except json.JSONDecodeError:
            continue
        if isinstance(items, list):
            for topic in items:
                topic_str = (topic or "").strip()
                if topic_str:
                    topic_set.add(topic_str)
    conn.close()
    topics = sorted(topic_set, key=lambda val: val.lower())
    return {"submitters": submitters, "assignees": assignees, "topics": topics}


@app.get("/api/questions/search")
def search_questions(submitter: str = "", assignee: str = "", topic: str = ""):
    submitter = (submitter or "").strip()
    assignee = (assignee or "").strip()
    topic_filter = (topic or "").strip().lower()
    conn = get_db()
    cur = conn.cursor()
    query = """
        SELECT
            q.*,
            m.meeting_date AS _meeting_date,
            m.commission_name AS _commission_name
        FROM questions q
        JOIN meetings m ON m.id = q.meeting_id
    """
    conditions = []
    params = []
    if submitter:
        value = f"%{submitter.lower()}%"
        conditions.append(
            """
            LOWER(TRIM(COALESCE(q.submitter_given_name, '') || ' ' || COALESCE(q.submitter_family_name, ''))) LIKE ?
            """
        )
        params.append(value)
    if assignee:
        value = f"%{assignee.lower()}%"
        conditions.append(
            """
            (
              LOWER(COALESCE(q.assignee_label, '')) LIKE ?
              OR LOWER(TRIM(COALESCE(q.assignee_given_name, '') || ' ' || COALESCE(q.assignee_family_name, ''))) LIKE ?
            )
            """
        )
        params.extend([value, value])
    if conditions:
        query += " WHERE " + " AND ".join(conditions)
    query += " ORDER BY m.meeting_date DESC, q.sequence_nr"

    cur.execute(query, params)
    rows = cur.fetchall()
    results = []
    for row in rows:
        item = _deserialize_question_row(row)
        topics_list = item.get("topics") or []
        if topic_filter:
            match = any(topic_filter in (t or "").lower() for t in topics_list)
            if not match:
                continue
        item["meeting_date"] = row["_meeting_date"]
        item["commission_name"] = row["_commission_name"]
        submitter_full = (
            f"{(item.get('submitter_given_name') or '').strip()} {(item.get('submitter_family_name') or '').strip()}".strip()
        )
        item["submitter_full_name"] = submitter_full or ""
        assignee_full = (
            (item.get("assignee_label") or "").strip()
            or f"{(item.get('assignee_given_name') or '').strip()} {(item.get('assignee_family_name') or '').strip()}".strip()
        )
        item["assignee_full_name"] = assignee_full or ""
        item["question_url"] = f"/meeting/{item['meeting_id']}#question-{item['id']}"
        results.append(item)

    conn.close()
    return {"questions": results}


@app.post("/api/councillors", status_code=201)
def create_councillor(payload: CouncillorCreate):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute(
            """INSERT INTO councillors (given_name, family_name, name_with_title, wrong_spellings)
               VALUES (?, ?, ?, ?)""",
            (
                payload.given_name.strip(),
                payload.family_name.strip(),
                (payload.name_with_title or "").strip(),
                (payload.wrong_spellings or "").strip(),
            ),
        )
        conn.commit()
        councillor_id = cur.lastrowid
        cur.execute("SELECT * FROM councillors WHERE id = ?", (councillor_id,))
        item = dict(cur.fetchone())
        return item
    except sqlite3.IntegrityError:
        conn.rollback()
        return JSONResponse(
            {"error": "Er bestaat al een raadslid met deze voor- en achternaam."},
            status_code=400,
        )
    finally:
        conn.close()


@app.patch("/api/councillors/{councillor_id}")
def update_councillor(councillor_id: int, payload: CouncillorUpdate):
    data = {k: v for k, v in payload.model_dump(exclude_unset=True).items()}
    if not data:
        return {"status": "no changes"}

    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT id FROM councillors WHERE id = ?", (councillor_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        return JSONResponse({"error": "councillor not found"}, status_code=404)

    fields = []
    values = []
    for k, v in data.items():
        fields.append(f"{k} = ?")
        values.append((v or "").strip())

    values.append(councillor_id)
    try:
        cur.execute(f"UPDATE councillors SET {', '.join(fields)} WHERE id = ?", values)
        conn.commit()
        cur.execute("SELECT * FROM councillors WHERE id = ?", (councillor_id,))
        item = dict(cur.fetchone())
        return item
    except sqlite3.IntegrityError:
        conn.rollback()
        return JSONResponse(
            {"error": "Er bestaat al een raadslid met deze voor- en achternaam."},
            status_code=400,
        )
    finally:
        conn.close()


@app.delete("/api/councillors/{councillor_id}")
def delete_councillor(councillor_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("DELETE FROM councillors WHERE id = ?", (councillor_id,))
    if cur.rowcount == 0:
        conn.close()
        return JSONResponse({"error": "councillor not found"}, status_code=404)
    conn.commit()
    conn.close()
    return {"status": "deleted", "councillor_id": councillor_id}
